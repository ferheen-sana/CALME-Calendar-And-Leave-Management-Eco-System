from django.http import HttpResponse
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from datetime import date


def get_designation(user):
    try:
        return user.profile.designation
    except:
        return "—"


# =========================
# EXCEL EXPORT
# =========================
def export_excel(leaves, title, user):
    wb = Workbook()
    ws = wb.active
    ws.title = "Leave Report"

    ws.append(["", title])
    ws.append(["Generated On", date.today().strftime("%d-%m-%Y")])
    ws.append([])

    headers = [
        "From",
        "Designation",
        "To",
        "Subject",
        "Leave Period",
        "Total Days",
        "Status",
        "Employee Signature",
        "Manager Signature",
    ]
    ws.append(headers)

    for l in leaves:
        ws.append([
            l.user.get_full_name() or l.user.username,
            get_designation(l.user),
            "Manager",
            "Leave Application",
            f"{l.start_date} to {l.end_date}",
            l.total_days,
            l.status,
            "",
            "",
        ])

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = 'attachment; filename="leave_report.xlsx"'
    wb.save(response)
    return response


# =========================
# PDF EXPORT
# =========================
def export_pdf(leaves, title, user):
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="leave_report.pdf"'

    p = canvas.Canvas(response, pagesize=A4)
    width, height = A4

    y = height - 50
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, y, title)

    y -= 30
    p.setFont("Helvetica", 10)
    p.drawString(50, y, f"Generated on: {date.today().strftime('%d-%m-%Y')}")

    y -= 40

    for l in leaves:
        text = p.beginText(50, y)
        text.textLine(f"Employee: {l.user.get_full_name() or l.user.username}")
        text.textLine(f"Designation: {get_designation(l.user)}")
        text.textLine(f"Leave Period: {l.start_date} to {l.end_date}")
        text.textLine(f"Total Days: {l.total_days}")
        text.textLine(f"Status: {l.status}")
        text.textLine("")
        text.textLine("Employee Signature: ________________")
        text.textLine("Manager Signature: ________________")
        p.drawText(text)

        y -= 140
        if y < 100:
            p.showPage()
            y = height - 50

    p.save()
    return response


# =========================
# WORD EXPORT
# =========================
def export_word(leaves, title, user):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated on: {date.today().strftime('%d-%m-%Y')}")

    for l in leaves:
        doc.add_paragraph(f"Employee: {l.user.get_full_name() or l.user.username}")
        doc.add_paragraph(f"Designation: {get_designation(l.user)}")
        doc.add_paragraph(f"Leave Period: {l.start_date} to {l.end_date}")
        doc.add_paragraph(f"Total Days: {l.total_days}")
        doc.add_paragraph(f"Status: {l.status}")
        doc.add_paragraph("\nEmployee Signature: ________________")
        doc.add_paragraph("Manager Signature: ________________")
        doc.add_page_break()

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="leave_report.docx"'
    doc.save(response)
    return response
