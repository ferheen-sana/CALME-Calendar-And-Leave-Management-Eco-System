from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from datetime import date, timedelta
import calendar
import json
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
import calendar, datetime
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from .models import Leave
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from datetime import datetime, timedelta
from accounts.models import Profile
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib import messages
from datetime import date
from collections import defaultdict
from django.http import JsonResponse
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from docx import Document
from django.contrib.auth.decorators import login_required
from django.shortcuts import redirect
from .utils_export import export_excel, export_pdf, export_word
from .models import Leave
from calendar import monthrange
from datetime import datetime, timedelta, time
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from reportlab.platypus import SimpleDocTemplate, Table
from reportlab.lib.pagesizes import A4
from openpyxl import Workbook
from .models import Meeting
from django.utils.timezone import now
from django.utils.dateparse import parse_datetime
from .models import Notification
from .models import Leave, LeaveBalance
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from .models import Leave
from django.views.decorators.csrf import csrf_exempt
from django.utils.dateparse import parse_datetime
from django.utils.timezone import make_aware
import openpyxl
from django.db.models import Sum, F, ExpressionWrapper, IntegerField
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Pt
from datetime import date
from openpyxl import Workbook
from django.http import HttpResponse
from .models import LeaveFileCounter
from django.conf import settings
import os
from reportlab.platypus import Image

from .models import Leave
from .utils import ADA_HOLIDAYS, RH_HOLIDAYS

@login_required
def dashboard(request):
    # ✅ SHOW PENDING + APPROVED IMMEDIATELY
    leaves = Leave.objects.exclude(status='REJECTED')

    calendar_data = defaultdict(list)

    for leave in leaves:
        profile = Profile.objects.get(user=leave.user)

        current = leave.start_date
        while current <= leave.end_date:
            calendar_data[str(current)].append({
                "name": leave.user.username,
                "color": profile.color
            })
            current += timedelta(days=1)

    return render(request, "dashboard.html", {
        "calendar_data": dict(calendar_data)
    })


@login_required
def apply_leave(request):
    if request.method == 'POST':
        Leave.objects.create(
            user=request.user,
            start_date=request.POST.get('from_date'),
            end_date=request.POST.get('to_date'),
            leave_type=request.POST.get('leave_type', 'CL'),  # ✅ safe default
            total_days=request.POST.get('total_days', 1),
            reason=request.POST.get('reason', ''),
            status='PENDING'   # shows immediately
        )
        return redirect('dashboard')

    return render(request, 'apply_leave.html')


@login_required
def calendar_view(request):
    year = date.today().year
    data = {}

    leaves = Leave.objects.filter(
        start_date__year=year,
        status="APPROVED"
    )

    for m in range(1, 13):
        month_name = calendar.month_name[m]
        data[month_name] = []

        for d in range(1, calendar.monthrange(year, m)[1] + 1):
            day_date = date(year, m, d)

            day_leaves = leaves.filter(
                start_date__lte=day_date,
                end_date__gte=day_date
            )

            if day_leaves.count() > 1:
                color = "red"
            elif day_leaves.count() == 1:
                color = day_leaves[0].user.profile.color
            else:
                color = "#e5e7eb"

            tooltip = ", ".join(l.user.username for l in day_leaves)

            data[month_name].append({
                "date": d,
                "color": color,
                "tooltip": tooltip
            })

    return render(request, "calendar.html", {
        "calendar": data,
        "year": year
    })

@login_required
def apply_leave(request):
    if request.method == 'POST':
        Leave.objects.create(
            user=request.user,
            start_date=request.POST['from_date'],
            end_date=request.POST['to_date'],
            leave_type=request.POST['leave_type'],
            total_days=request.POST.get('total_days', 1),
            reason=request.POST.get('reason', ''),
            status='PENDING'   # ✅ shows immediately
        )
    year = date.today().year
    months = []

    for m in range(1, 13):
        months.append({
            'name': calendar.month_name[m],
            'days': calendar.monthrange(year, m)[1],
            'month': m
        })

    if request.method == 'POST':
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        leave_type = request.POST.get('leave_type')
        session = request.POST.get('session')
        total_days = request.POST.get('total_days')

        Leave.objects.create(
            user=request.user,
            start_date=start_date,
            end_date=end_date,
            leave_type=leave_type,
            session=session,
            total_days=total_days,
            status='PENDING'
        )
        return redirect('/history/')

    return render(request, 'apply_leave.html', {
        'months': months,
        'year': year
    })


@login_required
def leave_history(request):
    leaves = Leave.objects.filter(user=request.user)

    years = (
        Leave.objects
        .filter(user=request.user)
        .dates("start_date", "year")
    )
    years = [y.year for y in years]

    months = [
        (1,"January"),(2,"February"),(3,"March"),(4,"April"),
        (5,"May"),(6,"June"),(7,"July"),(8,"August"),
        (9,"September"),(10,"October"),(11,"November"),(12,"December")
    ]

    return render(request, "leave_history.html", {
        "leaves": leaves,
        "years": years,
        "months": months,
        "employee_name": request.user.get_full_name(),
        "designation": getattr(request.user.profile, "designation", ""),
        "manager_name": "Manager"   # ✅ SAFE DEFAULT
    })


@login_required
def manager_dashboard(request):
    leaves = Leave.objects.filter(status="PENDING")
    return render(request,"manager_dashboard.html",{"leaves":leaves})


@login_required
def approve_leave(request,id):
    leave = get_object_or_404(Leave,id=id)
    leave.status="APPROVED"
    leave.save()
    return redirect("manager_dashboard")


@login_required
def reject_leave(request,id):
    leave = get_object_or_404(Leave,id=id)
    leave.status="REJECTED"
    leave.save()
    return redirect("manager_dashboard")

@login_required
def calendar_view(request):
    year = date.today().year
    data = {}

    leaves = Leave.objects.filter(
        start_date__isnull=False,
        end_date__isnull=False,
        start_date__year=year,
        status="APPROVED"
    )

    for m in range(1, 13):
        month_name = calendar.month_name[m]
        data[month_name] = []

        total_days = calendar.monthrange(year, m)[1]

        for d in range(1, total_days + 1):
            day_date = date(year, m, d)

            # HOLIDAYS FIRST
            if str(day_date) in ADA_HOLIDAYS:
                color = "green"
                tooltip = "ADA Holiday"

            elif str(day_date) in RH_HOLIDAYS:
                color = "yellow"
                tooltip = "RH Holiday"

            else:
                day_leaves = leaves.filter(
                    start_date__lte=day_date,
                    end_date__gte=day_date
                )

                if day_leaves.count() > 1:
                    color = "red"
                    tooltip = ", ".join(l.user.username for l in day_leaves)

                elif day_leaves.count() == 1:
                    color = day_leaves[0].user.profile.color
                    tooltip = day_leaves[0].user.username

                else:
                    color = "#e5e7eb"
                    tooltip = ""

            data[month_name].append({
                "date": d,
                "color": color,
                "tooltip": tooltip
            })

    return render(request, "calendar.html", {
        "calendar": data,
        "year": year
    })

def is_manager(user):
    return user.is_staff or user.is_superuser

@login_required
@user_passes_test(is_manager)
def manage_leaves(request):
    leaves = Leave.objects.filter(status='PENDING').order_by('applied_on')
    return render(request, 'manage_leaves.html', {'leaves': leaves})

@login_required
def notifications_page(request):
    # Pending leaves = notifications
    notifications = Leave.objects.filter(
        status="PENDING"
    ).order_by("-applied_on")

    notification_count = notifications.count()

    return render(
        request,
        "notifications.html",
        {
            "notifications": notifications,
            "notification_count": notification_count,
        }
    )

@login_required
def realtime_notifications(request):
    leaves = Leave.objects.filter(
        user=request.user,
        status__in=["APPROVED", "REJECTED"],
        notified=False
    )

    data = []
    for l in leaves:
        data.append({
            "id": l.id,
            "status": l.status,
            "from": l.start_date.strftime("%d %b %Y"),
            "to": l.end_date.strftime("%d %b %Y"),
            "days": l.total_days
        })
        l.notified = True   # 🔒 mark as delivered
        l.save()

    return JsonResponse({"count": len(data), "items": data})


@login_required
def leave_calendar(request):
    year = int(request.GET.get('year', date.today().year))

    leaves = Leave.objects.filter(
        status='APPROVED',
        start_date__year=year
    )

    calendar_data = defaultdict(list)

    for leave in leaves:
        d = leave.start_date
        while d <= leave.end_date:
            calendar_data[str(d)].append({
                "name": leave.user.username,
                "color": leave.user.profile.color
            })
            d = d.replace(day=d.day + 1)

    context = {
        "year": year,
        "months": range(1, 13),
        "calendar": calendar_data,
        "ada_holidays": ADA_HOLIDAYS,
        "rh_holidays": RH_HOLIDAYS,
        "calendar_module": calendar,
    }
    return render(request, "calendar.html", context)

def is_holiday(d):
    return d.weekday() >= 5 or str(d) in ADA_HOLIDAYS or str(d) in RH_HOLIDAYS


@login_required
def apply_leave(request):
    if request.method == "POST":
        start = datetime.strptime(
            request.POST["from_date"], "%Y-%m-%d"
        ).date()
        end = datetime.strptime(
            request.POST["to_date"], "%Y-%m-%d"
        ).date()

        Leave.objects.create(
            user=request.user,
            start_date=start,
            end_date=end,
            leave_type=request.POST["leave_type"],  # ✅ CL / RH ONLY
            total_days=float(request.POST["total_days"]),
            reason=request.POST.get("reason"),
            status="PENDING"
        )

        messages.success(request, "Leave applied successfully")
        return redirect("/calendar/")

    return render(request, "apply_leave.html")


@login_required
def export_employee_excel(request):
    leaves = Leave.objects.filter(user=request.user)
    return export_excel(leaves, "My Leave History")


@login_required
def export_employee_pdf(request):
    leaves = Leave.objects.filter(user=request.user)
    return export_pdf(leaves, "My Leave History")


@login_required
def export_manager_word(request):
    if not request.user.is_staff:
        return redirect("dashboard")
    leaves = Leave.objects.all()
    return export_word(leaves, "All Employees Leave Report")

@login_required
def export_employee_word(request):
    leaves = Leave.objects.filter(user=request.user)
    return export_word(leaves, "My Leave History")


@login_required
def export_manager_word(request):
    if not request.user.is_staff:
        return redirect("dashboard")
    leaves = Leave.objects.all()
    return export_word(leaves, "All Employees Leave Report")

@login_required
def export_manager_pdf(request):
    if not request.user.is_staff:
        return redirect("dashboard")
    leaves = Leave.objects.all()
    return export_pdf(leaves, "All Employees Leave Report")

@login_required
def export_manager_excel(request):
    if not request.user.is_staff:
        return redirect("dashboard")
    leaves = Leave.objects.all()
    return export_excel(leaves, "All Employees Leave Report", request.user)

def login_view(request):
    if request.method == "POST":
        user = authenticate(
            username=request.POST["username"],
            password=request.POST["password"]
        )

        if user:
            login(request, user)

            # ✅ ROLE BASED REDIRECT (FIX)
            if user.profile.role_type == 'MAN':
                return redirect('/manager/')
            else:
                return redirect('/')

        messages.error(request, "Invalid username or password")

    return render(request, "auth/login.html")


def register_view(request):
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        role_type = request.POST["role_type"]
        designation = request.POST["designation"]

        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already exists")
        else:
            user = User.objects.create_user(
                username=username,
                password=password
            )

            # ✅ PROFILE IS ALREADY CREATED BY SIGNAL
            profile = user.profile
            profile.role_type = role_type
            profile.designation = designation
            profile.save()

            messages.success(request, "Registration successful")
            return redirect("/login/")

    return render(request, "auth/register.html")


def logout_view(request):
    logout(request)
    return redirect("/login/")

@login_required
def employee_dashboard(request):
    if request.user.profile.role_type == 'MAN':
        return redirect('/manager/')
    return render(request, 'dashboard.html')


@login_required
def manager_dashboard(request):
    if request.user.profile.role_type != 'MAN':
        return redirect('/')
    return render(request, 'manager_dashboard.html')

@login_required
def delete_leave(request, leave_id):
    leave = get_object_or_404(Leave, id=leave_id, user=request.user)

    # ❗ Only pending leave can be deleted
    if leave.status == "PENDING":
        leave.delete()

    return redirect("leave_history")

def export_excel(leaves, title):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Leave Report"

    ws.append([title, ""])
    ws.append(["Generated On", date.today().strftime("%d %B %Y")])
    ws.append(["", ""])

    for leave in leaves:
        try:
            profile = Profile.objects.get(user=leave.user)
            designation = profile.designation
        except:
            designation = ""

        ws.append(("From", leave.user.get_full_name() or leave.user.username))
        ws.append(("Designation", designation))
        ws.append(("To", "Reporting Manager"))
        ws.append(("Subject", "Leave Application"))
        ws.append((
            "Leave Period",
            f"{leave.start_date} to {leave.end_date}"
        ))
        ws.append(("Total Days", leave.total_days))
        ws.append(("Status", leave.status))
        ws.append(("Employee Signature", ""))
        ws.append(("Manager Signature", ""))
        ws.append(("Date", date.today().strftime("%d %B %Y")))
        ws.append(("", ""))

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 50

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="{title}.xlsx"'
    wb.save(response)
    return response


def export_pdf(leaves, title):
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{title}.pdf"'

    c = canvas.Canvas(response, pagesize=A4)
    width, height = A4
    y = height - 60

    for leave in leaves:
        try:
            profile = Profile.objects.get(user=leave.user)
            designation = profile.designation
        except:
            designation = ""

        c.setFont("Helvetica-Bold", 15)
        c.drawCentredString(width / 2, y, "LEAVE APPLICATION")
        y -= 40

        c.setFont("Helvetica", 11)
        c.drawString(50, y, f"Date: {date.today().strftime('%d %B %Y')}")
        y -= 30

        c.drawString(50, y, "From,")
        y -= 15
        c.drawString(50, y, leave.user.get_full_name() or leave.user.username)
        y -= 15
        c.drawString(50, y, designation)
        y -= 30

        c.drawString(50, y, "To,")
        y -= 15
        c.drawString(50, y, "Reporting Manager")
        y -= 30

        c.setFont("Helvetica-Bold", 11)
        c.drawString(50, y, "Subject: Leave Application")
        y -= 30

        c.setFont("Helvetica", 11)
        text = c.beginText(50, y)
        text.textLine("Respected Sir/Madam,")
        text.textLine("")
        text.textLine(
            f"I would like to inform you that I have applied leave from "
            f"{leave.start_date} to {leave.end_date} "
            f"for {leave.total_days} day(s)."
        )
        text.textLine("")
        text.textLine("Kindly grant approval.")
        text.textLine("")
        text.textLine("Thanking you.")
        c.drawText(text)

        y = text.getY() - 50

        c.drawString(50, y, "Yours sincerely,")
        y -= 50
        c.drawString(50, y, "_________________________")
        c.drawString(50, y - 15, "Employee Signature")

        c.drawString(350, y, "_________________________")
        c.drawString(350, y - 15, "Manager Signature")

        c.showPage()
        y = height - 60

    c.save()
    return response

def export_word(leaves, title):
    doc = Document()

    for leave in leaves:
        employee_name = leave.user.get_full_name() or leave.user.username
        designation = get_designation(leave.user)
        today = date.today().strftime("%d %B %Y")

        # ---------- TITLE ----------
        heading = doc.add_heading("LEAVE APPLICATION", level=1)
        heading.alignment = 1  # Center

        doc.add_paragraph(f"Date: {today}")
        doc.add_paragraph("")

        # ---------- FROM ----------
        doc.add_paragraph("From,")
        doc.add_paragraph(employee_name)
        doc.add_paragraph(designation)
        doc.add_paragraph("")

        # ---------- TO ----------
        doc.add_paragraph("To,")
        doc.add_paragraph("Reporting Manager")
        doc.add_paragraph("")

        # ---------- SUBJECT ----------
        p = doc.add_paragraph()
        p.add_run("Subject: ").bold = True
        p.add_run("Leave Application")
        doc.add_paragraph("")

        # ---------- BODY ----------
        doc.add_paragraph("Respected Sir/Madam,")
        doc.add_paragraph(
            f"I would like to inform you that I have applied leave "
            f"from {leave.start_date} to {leave.end_date} "
            f"for {leave.total_days} day(s)."
        )
        doc.add_paragraph("Kindly grant approval.")
        doc.add_paragraph("")
        doc.add_paragraph("Thanking you.")
        doc.add_paragraph("")
        doc.add_paragraph("Yours sincerely,")
        doc.add_paragraph(employee_name)
        doc.add_paragraph("")

        # ---------- SIGNATURE TABLE ----------
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"

        table.cell(0, 0).text = "Employee Signature"
        table.cell(0, 1).text = "Manager Signature"

        table.cell(1, 0).text = "\n\n\nDate: __________"
        table.cell(1, 1).text = "\n\n\nDate: __________"

        doc.add_page_break()

    # ---------- RESPONSE ----------
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{title}.docx"'
    doc.save(response)
    return response

def get_designation(user):
    try:
        return user.profile.designation
    except:
        return "—"
    

@login_required
def manager_dashboard(request):
    # manager sees ALL approved + pending leaves
    leaves = Leave.objects.exclude(status='REJECTED')

    calendar_data = defaultdict(list)

    for leave in leaves:
        profile = Profile.objects.get(user=leave.user)

        current = leave.start_date
        while current <= leave.end_date:
            calendar_data[str(current)].append({
                "name": leave.user.username,
                "color": profile.color
            })
            current += timedelta(days=1)

    return render(request, "manager_dashboard.html", {
        "calendar_data": dict(calendar_data)
    })

@login_required
def manager_advanced_calendar(request):
    start = date.today() - timedelta(days=2)
    days = []

    for i in range(7):
        d = start + timedelta(days=i)

        leaves = Leave.objects.filter(
            start_date__lte=d,
            end_date__gte=d,
            status="APPROVED"
        )

        events = []
        for l in leaves:
            events.append({
                "user": l.user.username,
                "dept": l.user.profile.department or "General",
                "reason": l.reason,
                "color": l.user.profile.color or "#fde68a",
                "start": "09:00",
                "end": "18:00"
            })

        days.append({
            "date": d,
            "weekday": d.strftime("%A"),
            "events": events
        })

    return render(request, "manager/calendar_advanced.html", {
        "days": days
    })

def approve_manager(request):
    if request.method == "POST":
        leave_id = request.POST.get("leave_id")
        action = request.POST.get("action")
        leave_type = request.POST.get("leave_type")   # ✅ CL / RH


        leave = Leave.objects.get(id=leave_id)

        if action == "approve":
            leave.status = "APPROVED"
        elif action == "reject":
            leave.status = "REJECTED"

        leave.save()
        return redirect("leaves:approve_manager")  # ✅ FIXED

    leaves = Leave.objects.all().order_by("-applied_on")

    return render(request, "manager/approve_manager.html", {
        "leaves": leaves
    })

def is_manager(user):
    return user.profile.role_type == "MAN"

@login_required
def employee_history(request):
    leaves = Leave.objects.select_related("user", "user__profile")

    q = request.GET.get("q")
    month = request.GET.get("month")
    year = request.GET.get("year")

    if q:
        leaves = leaves.filter(user__username__icontains=q)
    if month:
        leaves = leaves.filter(start_date__month=month)
    if year:
        leaves = leaves.filter(start_date__year=year)

    years = Leave.objects.dates("start_date", "year")
    years = [y.year for y in years]

    months = [
        (1,"Jan"),(2,"Feb"),(3,"Mar"),(4,"Apr"),
        (5,"May"),(6,"Jun"),(7,"Jul"),(8,"Aug"),
        (9,"Sep"),(10,"Oct"),(11,"Nov"),(12,"Dec")
    ]

    return render(request, "manager/employee_history.html", {
        "leaves": leaves.order_by("-start_date"),
        "months": months,
        "years": years,
    })

@login_required
def export_employee_excel(request):
    return HttpResponse("Employee Excel Export")


@login_required
def export_employee_pdf(request):
    return HttpResponse("Employee PDF Export")


@login_required
def export_employee_word(request):
    return HttpResponse("Employee Word Export")


@login_required
def export_manager_excel(request):
    wb = Workbook()
    ws = wb.active
    ws.title = "Employee Leave History"

    ws.append([
        "Employee",
        "Leave Type",
        "From",
        "To",
        "Reason",
        "Status"
    ])

    for leave in Leave.objects.all():
        ws.append([
            leave.user.username,
            leave.leave_type,
            leave.start_date,
            leave.end_date,
            leave.reason,
            leave.status
        ])

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = "attachment; filename=employee_leaves.xlsx"

    wb.save(response)
    return response# ✅ FIXED

@login_required
def export_manager_pdf(request, leave_id):
    leave = get_object_or_404(Leave, id=leave_id)

    # ---------- SAFE DATA ----------
    employee_name = leave.user.get_full_name().strip()
    if not employee_name:
        employee_name = getattr(leave.user.profile, "full_name", "").strip()
    if not employee_name:
        employee_name = leave.user.username

    designation = getattr(leave.user.profile, "designation", "").strip() or "Employee"

    officer_name = "Reporting manager"
    officer_designation = "Project Director"

    division_name = "LCA Mk1 Division"

    company_name = "AERONAUTICAL DEVELOPMENT AGENCY"
    company_sub = "Ministry of Defence, Government of India"
    company_city = "Bengaluru – 560017"

    file_no = LeaveFileCounter.next_file_number()
    today_date = date.today().strftime("%d-%m-%Y")
    # --------------------------------

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = "attachment; filename=ADA_Leave_Letter.pdf"

    doc = SimpleDocTemplate(
        response,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()
    story = []

    # 🔹 HEADER
    story.append(Paragraph(
        f"<b>{company_name}</b><br/>"
        f"{company_sub}<br/>"
        f"{company_city}",
        styles["Title"]
    ))
    story.append(Spacer(1, 20))

    # 🔹 FILE NO + DATE (TABLE – OFFICIAL WAY)
    from reportlab.platypus import Table, TableStyle

    table = Table(
        [[f"File No: {file_no}", f"Date: {today_date}"]],
        colWidths=[260, 260]
    )
    table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(table)
    story.append(Spacer(1, 25))

    # 🔹 FROM
    story.append(Paragraph(
        f"From,<br/>"
        f"<b>{employee_name}</b><br/>"
        f"{designation}",
        styles["Normal"]
    ))
    story.append(Spacer(1, 15))

    # 🔹 TO
    story.append(Paragraph(
        f"To,<br/>"
        f"<b>{officer_name}</b><br/>"
        f"{division_name}",
        styles["Normal"]
    ))
    story.append(Spacer(1, 20))

    # 🔹 SUBJECT
    story.append(Paragraph(
        "<b>Subject: Grant of Leave</b>",
        styles["Normal"]
    ))
    story.append(Spacer(1, 20))

    # 🔹 BODY
    story.append(Paragraph(
        f"Sir/Madam,<br/><br/>"
        f"This is to certify that <b>{employee_name}</b>, {designation}, "
        f"is hereby granted <b>{leave.leave_type}</b> leave for a period of "
        f"<b>{leave.total_days}</b> day(s) from "
        f"{leave.start_date.strftime('%d %B %Y')} to "
        f"{leave.end_date.strftime('%d %B %Y')}.<br/><br/>"
        f"The leave is sanctioned as per CCS (Leave) Rules, 1972, "
        f"as applicable.",
        styles["Normal"]
    ))

    story.append(Spacer(1, 35))

    # 🔹 EMPLOYEE CLOSING
    story.append(Paragraph("Yours faithfully,<br/><br/>", styles["Normal"]))

    emp_sign = os.path.join(settings.BASE_DIR, "static", "employee_sign.png")
    if os.path.exists(emp_sign):
        story.append(Image(emp_sign, width=120, height=45))
        story.append(Spacer(1, 5))

    story.append(Paragraph(
        f"<b>{employee_name}</b><br/>{designation}",
        styles["Normal"]
    ))

    story.append(Spacer(1, 35))

    # 🔹 OFFICER APPROVAL
    story.append(Paragraph(
        "<b>Approved and Sanctioned,</b><br/><br/>",
        styles["Normal"]
    ))

    officer_sign = os.path.join(settings.BASE_DIR, "static", "signature.png")
    if os.path.exists(officer_sign):
        story.append(Image(officer_sign, width=120, height=45))
        story.append(Spacer(1, 5))

    story.append(Paragraph(
        f"<b>{officer_name}</b><br/>"
        f"{officer_designation}<br/>"
        f"{company_name}",
        styles["Normal"]
    ))

    doc.build(story)
    return response

@login_required
def export_manager_word(request, leave_id):
    leave = get_object_or_404(Leave, id=leave_id)

    # ---------- SAFE DATA (NO BLANKS EVER) ----------
    employee_name = leave.user.get_full_name().strip()
    if not employee_name:
        employee_name = getattr(leave.user.profile, "full_name", "").strip()
    if not employee_name:
        employee_name = leave.user.username

    designation = getattr(leave.user.profile, "designation", "").strip()
    if not designation:
        designation = "Employee"

    manager_name = "Reporting Manager"
    manager_designation = "LCA MK1"

    company_name = "Aeronautical Development Agency"
    company_city = "Bengaluru"

    file_no = LeaveFileCounter.next_file_number()
    today_date = date.today().strftime("%d-%m-%Y")
    # -----------------------------------------------

    doc = Document()

    # 🔹 COMPANY HEADER
    doc.add_heading(company_name, level=1)
    doc.add_paragraph(company_city)

    # 🔹 FILE NO + DATE (RIGHT FEEL)
    p = doc.add_paragraph()
    run = p.add_run(f"File No: {file_no}")
    run.bold = True
    p.add_run(f"\t\t\t\tDate: {today_date}")

    doc.add_paragraph("")  # spacing

    # 🔹 FROM
    doc.add_paragraph(
        f"From,\n"
        f"{employee_name}\n"
        f"{designation}\n"
    )

    # 🔹 TO
    doc.add_paragraph(
        f"To,\n"
        f"{manager_name}\n"
        f"{manager_designation}\n"
    )

    # 🔹 SUBJECT
    doc.add_paragraph("Subject: Leave Application\n")

    # 🔹 BODY (NAME WILL ALWAYS COME)
    doc.add_paragraph(
        f"Sir/Madam,\n\n"
        f"This is to certify that {employee_name}, {designation}, "
        f"is hereby granted {leave.leave_type} leave for a period of "
        f"{leave.total_days} day(s) from "
        f"{leave.start_date.strftime('%d %B %Y')} to "
        f"{leave.end_date.strftime('%d %B %Y')}. "
        f"The leave is sanctioned as per rules in force."
    )

    doc.add_paragraph("")  # spacing

    # 🔹 CLOSING – EMPLOYEE
    doc.add_paragraph("Yours faithfully,\n")

    doc.add_paragraph(
        f"{employee_name}\n"
        f"{designation}"
    )

    doc.add_paragraph("")  # spacing

    # 🔹 MANAGER APPROVAL
    doc.add_paragraph("Approved By,\n")

    doc.add_paragraph(
        f"{manager_name}\n"
        f"{manager_designation}"
    )

    # 🔹 RESPONSE
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = "attachment; filename=Leave_Letter.docx"

    doc.save(response)
    return response

def calendar_advanced(request):
    today = timezone.now().date()
    days = []

    for i in range(-2, 3):
        day = today + timedelta(days=i)
        days.append({
            "date": day
        })

    return render(
        request,
        "manager/calendar_advanced.html",
        {
            "days": days,
            "now": timezone.now()
        }
    )

@login_required
def manager_advanced_calendar(request):
    meetings = Meeting.objects.all().order_by("start")

    # auto-update status
    now = timezone.now()
    for m in meetings:
        if m.end < now:
            m.status = "COMPLETED"
        elif m.start > now:
            m.status = "UPCOMING"
        else:
            m.status = "PENDING"
        m.save(update_fields=["status"])

    return render(
        request,
        "manager/calendar_advanced.html",
        {"meetings": meetings}
    )


def add_meeting(request):
    if request.method == "POST":
        Meeting.objects.create(
            title=request.POST["title"],
            description=request.POST.get("description", ""),
            start=parse_datetime(request.POST["start"]),
            end=parse_datetime(request.POST["end"]),
        )
    return redirect("leaves:manager_advanced_calendar")


@login_required
def delete_meeting(request, meeting_id):
    meeting = get_object_or_404(Meeting, id=meeting_id)
    meeting.delete()
    return redirect("manager_advanced_calendar")

@login_required
def realtime_notifications(request):
    count = Leave.objects.filter(
        user=request.user,
        notified=False,
        status__in=["APPROVED", "REJECTED"]
    ).count()

    return JsonResponse({"count": count})

def notification_count(request):
    if not request.user.is_authenticated:
        return JsonResponse({"count": 0})

    count = Notification.objects.filter(
        user=request.user,
        is_read=False
    ).count()

    return JsonResponse({"count": count})

@login_required
def manager_notifications(request):
    pending_leaves = Leave.objects.filter(status="Pending").order_by("-id")
    if request.user.profile.role_type != "MAN":
        return redirect("leaves:dashboard")

    pending = Leave.objects.filter(status="PENDING")
    approved = Leave.objects.filter(status="APPROVED")
    rejected = Leave.objects.filter(status="REJECTED")

    return render(
        request,
        "manager/notifications.html",
        {
            "leaves": pending.select_related("user"),
            "pending_count": pending.count(),
            "approved_count": approved.count(),
            "rejected_count": rejected.count(),
            "pending_leaves": pending_leaves
        }
    )

@login_required
def apply_leave(request):

    # ✅ ensure balance exists
    balance, created = LeaveBalance.objects.get_or_create(
        user=request.user,
        defaults={"cl_balance": 8, "rh_balance": 2}
    )

    if request.method == "POST":
        leave_type = request.POST.get("leave_type")
        from_date = request.POST.get("from_date")
        to_date = request.POST.get("to_date")
        total_days = float(request.POST.get("total_days", 0))
        reason = request.POST.get("reason", "")

        # 🔴 VALIDATION
        if total_days <= 0:
            messages.error(request, "Invalid leave duration")
            return redirect("apply_leave")

        # 🔴 BALANCE CHECK
        if leave_type == "CL":
            if balance.cl_balance < total_days:
                messages.error(request, "Insufficient CL balance")
                return redirect("apply_leave")

            balance.cl_balance -= total_days

        elif leave_type == "RH":
            if balance.rh_balance < total_days:
                messages.error(request, "Insufficient RH balance")
                return redirect("apply_leave")

            balance.rh_balance -= total_days

        else:
            messages.error(request, "Invalid leave type")
            return redirect("apply_leave")

        # ✅ SAVE BALANCE
        balance.save()

        # ✅ SAVE LEAVE
        Leave.objects.create(
            user=request.user,
            start_date=from_date,
            end_date=to_date,
            leave_type=leave_type,
            total_days=total_days,
            reason=reason,
            status="PENDING"
        )

        messages.success(request, "Leave applied successfully")
        return redirect("dashboard")

    # 🔵 GET request → show updated balance
    context = {
        "cl_balance": balance.cl_balance,
        "rh_balance": balance.rh_balance
    }

    return render(request, "apply_leave.html", context)

@login_required
def delete_leave(request, leave_id):
    leave = get_object_or_404(
        Leave,
        id=leave_id,
        user=request.user
    )

    if leave.status == "PENDING":
        leave.delete()

    return redirect("leaves:leave_history")

def employee_history(request):
    leaves = Leave.objects.filter(user=request.user)

    context = {
        "leaves": leaves,
        "employee_name": request.user.get_full_name() or request.user.username,
        "designation": getattr(request.user, "designation", "Employee"),
        "manager_name": "Manager",
    }

    return render(request, "employee_history.html")  # ✅ your template

@login_required
def delete_leave(request, leave_id):
    Leave.objects.filter(id=leave_id, user=request.user).delete()
    return redirect("leaves:leave_history")



@login_required
def apply_leave(request):
    balance, _ = LeaveBalance.objects.get_or_create(user=request.user)

    if request.method == "POST":
        leave_type = request.POST["leave_type"]
        days = float(request.POST["total_days"])

        # ❌ insufficient balance
        if leave_type == "CL" and balance.cl_balance < days:
            messages.error(request, "Insufficient CL balance")
            return redirect("apply_leave")

        if leave_type == "RH" and balance.rh_balance < days:
            messages.error(request, "Insufficient RH balance")
            return redirect("apply_leave")

        # ✅ CREATE LEAVE  (❗ DO NOT SET total_days ❗)
        Leave.objects.create(
            user=request.user,
            leave_type=leave_type,
            start_date=request.POST["from_date"],
            end_date=request.POST["to_date"],
            reason=request.POST["reason"],
            status="PENDING",
        )

        # ✅ DEDUCT BALANCE
        if leave_type == "CL":
            balance.cl_balance -= days
        else:
            balance.rh_balance -= days

        balance.save()

        messages.success(request, "Leave applied successfully")
        return redirect("leaves:leave_history")


    return render(
        request,
        "apply_leave.html",
        {
            "cl_balance": balance.cl_balance,
            "rh_balance": balance.rh_balance,
        },
    )

@login_required
def delete_leave(request, leave_id):
    leave = get_object_or_404(Leave, id=leave_id, user=request.user)
    balance = LeaveBalance.objects.get(user=request.user)

    if leave.status == "PENDING":
        if leave.leave_type == "CL":
            balance.cl_balance += leave.total_days
        else:
            balance.rh_balance += leave.total_days
        balance.save()

    leave.delete()
    return redirect("leaves:leave_history")

@login_required
def reject_leave(request, leave_id):
    leave = Leave.objects.get(id=leave_id)
    balance = LeaveBalance.objects.get(user=leave.user)

    if leave.status != "REJECTED":
        if leave.leave_type == "CL":
            balance.cl_balance += leave.total_days
        else:
            balance.rh_balance += leave.total_days

        balance.save()
        leave.status = "REJECTED"
        leave.save()

    return redirect("approve_manager")
@login_required
def export_manager_excel(request):
    file_number = LeaveFileCounter.next_file_number()
    wb = Workbook()
    ws = wb.active
    ws.title = "Leave Report"

    ws.append([
        "File No", "Employee", "Designation",
        "Leave Type", "Start Date", "End Date",
        "Days", "Status"
    ])

    for leave in Leave.objects.select_related("user"):
        ws.append([
            leave.id,
            leave.user.get_full_name(),
            leave.user.profile.designation,
            leave.leave_type,
            leave.start_date,
            leave.end_date,
            leave.total_days,
            leave.status
        ])

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = "attachment; filename=leave_report.xlsx"

    wb.save(response)
    return response

def manager_notifications(request):
    leaves = Leave.objects.select_related('user').order_by('-applied_on')

    employees = User.objects.filter(leave__isnull=False).distinct()

    today = now().date()
    yesterday = today - timedelta(days=1)

    pending_count = leaves.filter(status="PENDING").count()
    approved_count = leaves.filter(status="APPROVED").count()
    rejected_count = leaves.filter(status="REJECTED").count()

    context = {
        "leaves": leaves,
        "employees": employees,
        "today": today,
        "yesterday": yesterday,
        "pending_count": pending_count,
        "approved_count": approved_count,
        "rejected_count": rejected_count,
    }
    return render(request, "manager/notifications.html", context)


def manager_notifications_json(request):
# 🔹 Pending count (ALWAYS correct)
    pending_count = Leave.objects.filter(status="PENDING").count()


# 🔹 Send ONLY newly applied leaves (last 1 minute)
    last_minute = now() - timedelta(minutes=1)


    new_leaves = (
        Leave.objects
        .select_related('user')
        .filter(status="PENDING", applied_on__gte=last_minute)
        .order_by('-applied_on')
        )


    leaves_list = []
    for l in new_leaves:
       leaves_list.append({
       "id": l.id,
       "user_id": l.user.id,
       "full_name": l.user.get_full_name() or l.user.username,
       "leave_type": l.leave_type,
       "start_date": l.start_date.strftime("%Y-%m-%d"),
       "end_date": l.end_date.strftime("%Y-%m-%d"),
       "reason": l.reason,
       "created_at": l.applied_on.isoformat(), # IMPORTANT for JS
    })


    return JsonResponse({
      "pending_count": pending_count,
      "leaves": leaves_list
    })

@csrf_exempt
def update_meeting(request):
    if request.method == "POST":
        meeting_id = request.POST.get("id")
        start = parse_datetime(request.POST.get("start"))
        end = parse_datetime(request.POST.get("end"))

        Meeting.objects.filter(id=meeting_id).update(
            start=start,
            end=end
        )
        return JsonResponse({"status":"ok"})
    
@login_required
@user_passes_test(is_manager)
def manager_advanced_calendar(request):
    meetings = Meeting.objects.all().order_by("start")

    meetings_json = [
        {
            "id": m.id,
            "title": m.title,
            "description": m.description or "",
            "start": m.start.isoformat(),
            "end": m.end.isoformat() if m.end else None,
            "repeat": m.repeat,
            "status": m.status(),
        }
        for m in meetings
    ]

    context = {
        "meetings_json": json.dumps(meetings_json),
    }
    return render(request, "manager/calendar_advanced.html", context)

def update_meeting(request):
    if request.method == "POST":
        m = Meeting.objects.get(id=request.POST["id"])
        m.start = make_aware(datetime.fromisoformat(request.POST["start"]))
        m.end = make_aware(datetime.fromisoformat(request.POST["end"]))
        m.save()
    return JsonResponse({"status": "ok"})

def add_meeting(request):
    m = Meeting.objects.create(
        title=request.POST["title"],
        start=parse_datetime(request.POST["start"]),
        end=parse_datetime(request.POST["end"]),
        repeat=request.POST.get("repeat", "none"),
        created_by=request.user
    )
    return JsonResponse({"id": m.id})

def edit_meeting(request, id):
    m = get_object_or_404(Meeting, id=id)
    m.title = request.POST["title"]
    m.start = parse_datetime(request.POST["start"])
    m.end = parse_datetime(request.POST["end"])
    m.repeat = request.POST.get("repeat", "none")
    m.save()
    return JsonResponse({"status": "updated"})

def delete_meeting(request, id):
    Meeting.objects.filter(id=id).delete()
    return JsonResponse({"status": "deleted"})

@csrf_exempt
def move_meeting(request):
    m = Meeting.objects.get(id=request.POST["id"])
    m.start = parse_datetime(request.POST["start"])
    m.end = parse_datetime(request.POST["end"])
    m.save()
    return JsonResponse({"status": "moved"})

@login_required
def manager_delete_leave(request, leave_id):
    leave = get_object_or_404(Leave, id=leave_id)
    leave.delete()
    return redirect("manager_dashboard")

def is_manager(user):
    return user.is_authenticated and user.is_staff


@login_required
@user_passes_test(is_manager)
def approve_leave(request, leave_id):
    leave = get_object_or_404(Leave, id=leave_id)

    if leave.status == "PENDING":
        leave.status = "APPROVED"
        leave.decision_on = timezone.now()
        leave.notified = False
        leave.save()

    return redirect("/manager/leaves/")   # ✅ HARD URL
    return redirect("leaves:approve_manager")



@login_required
@user_passes_test(is_manager)
def reject_leave(request, leave_id):
    leave = get_object_or_404(Leave, id=leave_id)

    if leave.status == "PENDING":
        leave.status = "REJECTED"
        leave.decision_on = timezone.now()
        leave.notified = False
        leave.save()

    return redirect("/manager/leaves/")   # ✅ HARD URL
 


# ===== MANAGER APPROVAL PAGE =====
@login_required
@user_passes_test(is_manager)
def approve_manager(request):
    leaves = Leave.objects.all().order_by("-applied_on")
    return render(request, "manager/approve_manager.html", {
        "leaves": leaves
    })

@login_required
def employee_history(request):

    leaves = Leave.objects.select_related("user").all()

    # ---------------- SEARCH ----------------
    q = request.GET.get("q")
    if q:
        leaves = leaves.filter(user__username__icontains=q)

    # ---------------- MONTH / YEAR FILTER ----------------
    month = request.GET.get("month")
    year = request.GET.get("year")

    if month:
        leaves = leaves.filter(start_date__month=int(month))

    if year:
        leaves = leaves.filter(start_date__year=int(year))

    # ---------------- DAY CALCULATION (ORM SAFE) ----------------
    duration = ExpressionWrapper(
        F("end_date") - F("start_date") + 1,
        output_field=IntegerField()
    )

    approved = leaves.filter(status="APPROVED")

    cl_total = approved.filter(
        leave_type="CL"
    ).aggregate(total=Sum(duration))["total"] or 0

    rh_total = approved.filter(
        leave_type="RH"
    ).aggregate(total=Sum(duration))["total"] or 0

    # ---------------- MONTHS & YEARS ----------------
    months = [
        (1, "Jan"), (2, "Feb"), (3, "Mar"), (4, "Apr"),
        (5, "May"), (6, "Jun"), (7, "Jul"), (8, "Aug"),
        (9, "Sep"), (10, "Oct"), (11, "Nov"), (12, "Dec")
    ]

    years = range(datetime.now().year - 5, datetime.now().year + 1)

    context = {
        "leaves": leaves.order_by("-start_date"),
        "cl_total": cl_total,
        "rh_total": rh_total,
        "months": months,
        "years": years,
    }

    return render(request, "manager/employee_history.html", context)

# ===== APPROVE LEAVE =====
def manager_advanced_calendar(request):
    meetings = Meeting.objects.filter(created_by=request.user)

    meetings_json = [
        {
            "id": m.id,
            "title": m.title,
            "start": m.start.isoformat(),
            "end": m.end.isoformat() if m.end else None,
        }
        for m in meetings
    ]

    return render(
        request,
        "manager/calendar_advanced.html",
        {"meetings_json": json.dumps(meetings_json)}
    )

def manager_advanced_calendar(request):
    events = []

    for m in Meeting.objects.all():
        events.append({
            "title": m.title,
            "start": m.start.isoformat(),
            "end": m.end.isoformat(),
        })

    return render(request, "manager/calendar_advanced.html", {
        "events": events
    })

def generate_file_number():
    current_year = date.today().year

    obj, created = LeaveFileCounter.objects.get_or_create(
        year=current_year,
        defaults={"last_number": 0}
    )

    obj.last_number += 1
    obj.save(update_fields=["last_number"])

    return f"LMS/HR/{current_year}/{obj.last_number:04d}"

def manager_notifications_count(request):
    pending_count = Leave.objects.filter(status="PENDING").count()
    return JsonResponse({"pending_count": pending_count})

@login_required
def employee_notifications_count(request):
    """
    Return JSON with:
    - pending_count: number of new notifications (approved/rejected leaves)
    - notifications: list of recent notifications
    """
    user = request.user

    # Get leaves that were updated after last seen (or last 24h)
    notifications = Leave.objects.filter(user=user).exclude(status='PENDING').order_by('-updated_at')[:5]

    pending_count = notifications.count()

    notification_list = []
    for leave in notifications:
        notification_list.append({
            'id': leave.id,
            'leave_type': leave.leave_type,
            'start_date': leave.start_date.strftime("%Y-%m-%d"),
            'end_date': leave.end_date.strftime("%Y-%m-%d"),
            'status': leave.status,  # APPROVED or REJECTED
        })

    return JsonResponse({
        'pending_count': pending_count,
        'notifications': notification_list,
    })